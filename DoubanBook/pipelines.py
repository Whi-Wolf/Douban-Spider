# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: https://docs.scrapy.org/en/latest/topics/item-pipeline.html


# useful for handling different item types with a single interface
from itemadapter import ItemAdapter
from docx import Document
from docx.oxml.ns import qn
from scrapy.pipelines.images import ImagesPipeline
from scrapy.exceptions import DropItem
from scrapy import Request



class DoubanbookPipeline:
    document=Document()
    def __init__(self):
        self.document = Document()
        
    
    #把书名作者和图片放进word文档，图文同一批，对应起来
    def process_item(self, item, spider):
        # # with open('book.txt','a',encoding='utf8')as fp:
        # #     fp.write(item['bookname']+'\t')
        # #     fp.write(item['author']+'\n')
        
        #设置英文+中文字体
        self.document.styles['Normal'].font.name = 'Times New Roman'
        self.document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        #先放书名+作者
        self.document.add_paragraph(item['bookname']+'\t'+item['author'])
        #把下载到路径里的图片放进word
        self.document.add_picture('.\Img\\'+item['bookname']+'.jpg')
        self.document.save('book.doc') #保存
        # remove(item['bookname']+'.jpg')
        return item
    
    

class BookImgDownloadPipeline(ImagesPipeline):
 
    # 设置下载文件请求的请求头
    # default_headers = {
    #     'accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
    #     'accept-encoding': 'gzip, deflate, br',
    #     'accept-language': 'zh-CN,zh;q=0.9',
    #     'referer': 'None',#'https://book.douban.com/',
    #     'user-agent': 'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/65.0.3314.0 Safari/537.36 SE 2.X MetaSr 1.0',
    # }
    #headers=self.default_headers
 
    # 伪装成站内请求，反反爬
    def get_media_requests(self, item, info):
        return Request(item['image_urls'])
        
        
 
    # 自定义 文件名(因为报错尝试了很多方法，结果是setting里的问题)
    def file_path(self, request, item, response=None, info=None):
        #path = '%s.jpg' % (item['bookname'])
        name = item['bookname']
        #print('******the results is********:',name)
        pic_name=name+'.jpg'
        return pic_name
        #return'full/%s.jpg' % (name)
        # img_name = request.url.split('/')[-1]
        # return img_name
        #image_guid = hashlib.sha1(to_bytes(request.url)).hexdigest()
        #return 'full/%s.jpg' % (image_guid)
 
    def item_completed(self, results, item, info):
        #print('******the results is********:', results) #也可以在这个方法中更改文件名，因为在file_path存储完成后，item_completed里有个results参数，results参数保存了图片下载的相关信息
        # image_paths = [x['path'] for ok, x in results if ok]
        # newname=item['bookname']+'.jpg'
        # os.rename("/neteaseauto/" + image_paths[0], "/neteaseauto/" + newname)
        return item

